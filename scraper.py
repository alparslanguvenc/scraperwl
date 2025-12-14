from playwright.sync_api import sync_playwright
from docx import Document
import time
import random

def scrape_reviews_playwright(url):
    reviews_data = []
    
    with sync_playwright() as p:
        # Launch browser (headless=True is default, but False helps debugging if needed)
        # Critical for Docker/Cloud: --no-sandbox
        browser = p.chromium.launch(
            headless=True,
            args=["--no-sandbox", "--disable-setuid-sandbox"]
        )
        page = browser.new_page(user_agent="Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.114 Safari/537.36")
        
        # Helper to ensure we are on the reviews page
        if "/reviews/" not in url:
            target_url = url.rstrip('/') + '/reviews/'
        else:
            target_url = url
            
        print(f"Navigating to {target_url}...")
        page.goto(target_url, timeout=60000)
        
        # Wait for initial load
        page.wait_for_load_state("networkidle")
        
        # LOCATORS (Based on previous inspection)
        # Content: div[class*="experienceReviewContent"]
        # Date: div[class*="experienceReviewDate"]
        # Title: span[class*="ReviewTitle"]
        
        # The user mentioned a "popup" or "modal".
        # If the reviews page is a standalone page (which /reviews/ seems to be), infinite scroll is likely on the main window.
        # If it's a modal on the main tour page, we'd handle that differently.
        # Given /reviews/ exists, we assume it's the full page list.
        
        # Scroll logic - Robust Infinite Scroll
        previous_count = 0
        scroll_attempts = 0
        max_scroll_attempts = 100 # Safety limit
        
        # Scroll logic - Optimized for "Load More" / Infinite Scroll
        previous_count = 0
        scroll_attempts = 0
        max_scroll_attempts = 150 # Increased limit
        
        # Locate the container. 
        # If it's an overlay, we need to scroll the overlay, not the window.
        # Try to find a scrollable container
        files = page.locator('div[class*="Overlay"]').all()
        scrollable_element = None
        if files:
            print("Found Overlay! attempting to scroll overlay.")
            scrollable_element = files[0]
        
        while scroll_attempts < max_scroll_attempts:
            # Method 1: Scroll to the last review item
            review_items = page.locator('div[class*="experienceReview_"]').all()
            if review_items:
                review_items[-1].scroll_into_view_if_needed()
            
            # Method 2: Press PageDown (Global)
            page.keyboard.press("PageDown")
            
            # Method 3: Scroll the specific container if found
            if scrollable_element:
                scrollable_element.evaluate("el => el.scrollTop = el.scrollHeight")
            else:
                page.evaluate("window.scrollTo(0, document.body.scrollHeight)")
            
            # Wait for potential network requests
            try: 
                page.wait_for_load_state("networkidle", timeout=3000) 
            except:
                pass 
                
            # Check if new items loaded
            current_count = page.locator('div[class*="experienceReviewContent"]').count()
            print(f"Scroll {scroll_attempts}: Found {current_count} reviews so far...")
            
            if current_count > previous_count:
                previous_count = current_count
                scroll_attempts = 0 # Reset attempts
                time.sleep(1.5) 
            else:
                scroll_attempts += 1
                time.sleep(1.5)
                
                if scroll_attempts > 3:
                    # Force a slightly random scroll to trigger events
                    page.evaluate("window.scrollBy(0, -300)")
                    time.sleep(0.5)
                    page.evaluate("window.scrollTo(0, document.body.scrollHeight)")
            
            if scroll_attempts >= 10: # Increased patience
                print("No new items loaded after 10 attempts. Stopping.")
                break
            
        print("Finished scrolling. Extracting data...")
        
        # Extract data using Playwright locators
        # Note: Class names might change, so we use partial match if possible or specific structure
        
        # We grab all review containers. 
        # Inspecting the HTML structure: 
        # <li class="_ListItem_132mq_19"> contains the review
        
        review_items = page.locator('li[class*="_ListItem_"]').all() # Rough guess based on "132mq" suffix seen in previous grep
        
        # Better strategy: Get all divs that look like content and find their blocks
        # Using query_selector_all for flexibility
        
        # Let's try to extract from the page content directly
        html = page.content()
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, 'html.parser')
        
        # Now use BS4 to parse the fully loaded HTML
        content_divs = soup.find_all('div', class_=lambda x: x and 'experienceReviewContent' in x)
        
        for div in content_divs:
            text = div.get_text(strip=True)
            
            title_span = div.find('span', class_=lambda x: x and 'ReviewTitle' in x)
            title = title_span.get_text(strip=True) if title_span else "No Title"
            
            # Find date in parent or siblings
            # Structure: 
            # <li class="ListItem...">
            #   <div ...> <figure Avatar> </figure> <div Details>... Date </div> </div>
            #   <div Content> ... </div>
            # </li>
            
            # Go up to finding the list item, then search down for date
            list_item = div.find_parent('li')
            date = "Unknown Date"
            if list_item:
                date_div = list_item.find('div', class_=lambda x: x and 'experienceReviewDate' in x)
                if date_div:
                    date = date_div.get_text(strip=True)
            
            reviews_data.append({
                'title': title,
                'text': text,
                'date': date
            })
            
        browser.close()
        
    print(f"Total reviews scraped: {len(reviews_data)}")
    return reviews_data

def save_to_docx(reviews, output_path):
    doc = Document()
    doc.add_heading('Withlocals Scraped Reviews', 0)
    
    if not reviews:
        doc.add_paragraph("No reviews found.")
    else:
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Date'
        hdr_cells[1].text = 'Title'
        hdr_cells[2].text = 'Review'
        
        for r in reviews:
            row_cells = table.add_row().cells
            row_cells[0].text = r['date']
            row_cells[1].text = r['title']
            row_cells[2].text = r['text']
            
    doc.save(output_path)
    print(f"Saved to {output_path}")
